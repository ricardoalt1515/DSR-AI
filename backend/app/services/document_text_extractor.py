"""Deterministic text extraction for office documents."""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO

MAX_EXTRACT_CHARS = 100000
MAX_XLSX_SHEETS = 20
MAX_XLSX_ROWS = 10000
MAX_XLSX_CELLS = 200000
MAX_DOCX_BLOCKS = 20000
TRUNCATION_MARKER = "[TRUNCATED: content shortened]"


@dataclass(frozen=True)
class ExtractedTextResult:
    text: str
    char_count: int
    truncated: bool


def extract_xlsx_text(file_bytes: bytes) -> ExtractedTextResult:
    load_workbook = _get_openpyxl_load_workbook()
    workbook = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    try:
        blocks: list[str] = []
        cell_count = 0
        row_count = 0
        for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
            if sheet_index > MAX_XLSX_SHEETS:
                break
            blocks.append(f"=== SHEET: {sheet.title} ===")
            sheet_row_count = 0
            for row in sheet.iter_rows(values_only=True):
                row_count += 1
                sheet_row_count += 1
                if row_count > MAX_XLSX_ROWS:
                    break
                values: list[str] = []
                for value in row:
                    cell_count += 1
                    if cell_count > MAX_XLSX_CELLS:
                        break
                    values.append(_normalize_text(_stringify_cell(value)))
                if cell_count > MAX_XLSX_CELLS:
                    break
                if values:
                    blocks.append("\t".join(values))
            if sheet_row_count == 0:
                blocks.pop()
            if row_count > MAX_XLSX_ROWS or cell_count > MAX_XLSX_CELLS:
                break
        return _finalize_blocks(blocks)
    finally:
        workbook.close()


def extract_docx_text(file_bytes: bytes) -> ExtractedTextResult:
    document = _load_docx_document(file_bytes)
    blocks: list[str] = []
    block_count = 0

    for paragraph in document.paragraphs:
        text = _normalize_text(paragraph.text)
        if text:
            blocks.append(text)
            block_count += 1
            if block_count >= MAX_DOCX_BLOCKS:
                return _finalize_blocks(blocks)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = _normalize_text(cell.text)
                if cell_text:
                    blocks.append(cell_text)
                    block_count += 1
                    if block_count >= MAX_DOCX_BLOCKS:
                        return _finalize_blocks(blocks)

    return _finalize_blocks(blocks)


def _get_openpyxl_load_workbook():
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("xlsx_parser_unavailable") from exc
    return load_workbook


def _load_docx_document(file_bytes: bytes):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("docx_parser_unavailable") from exc
    return Document(BytesIO(file_bytes))


def _stringify_cell(value: object) -> str:
    if value is None:
        return ""
    return str(value)


def _normalize_text(value: str) -> str:
    cleaned = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", value)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def _finalize_blocks(blocks: list[str]) -> ExtractedTextResult:
    text = "\n".join(blocks)
    text = text.strip()
    if not text:
        return ExtractedTextResult(text="", char_count=0, truncated=False)

    if len(text) <= MAX_EXTRACT_CHARS:
        return ExtractedTextResult(text=text, char_count=len(text), truncated=False)

    slice_len = max(0, MAX_EXTRACT_CHARS - len(TRUNCATION_MARKER))
    truncated_text = f"{text[:slice_len]}{TRUNCATION_MARKER}"
    return ExtractedTextResult(
        text=truncated_text,
        char_count=len(truncated_text),
        truncated=True,
    )
